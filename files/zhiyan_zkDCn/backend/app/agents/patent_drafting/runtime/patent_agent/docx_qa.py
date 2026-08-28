from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape
from zipfile import BadZipFile, ZipFile

from docx import Document
from PIL import Image
from pypdf import PdfReader

from patent_agent.errors import ExportError, ParseError


_CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff]")
_CJK_FONT_MARKERS = (
    "noto",
    "sourcehan",
    "pingfang",
    "hiragino",
    "songti",
    "heiti",
    "yahei",
    "simsun",
    "simhei",
    "arialunicode",
    "wenquanyi",
    "sarasa",
    "lxgw",
    "stsong",
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _normalized_font_name(name: str) -> str:
    name = re.sub(r"^[A-Z]{6}\+", "", name)
    return re.sub(r"[^a-z0-9]", "", name.casefold())


def _has_cjk_font_evidence(font_names: list[str]) -> bool:
    return any(
        marker in _normalized_font_name(font_name)
        for font_name in font_names
        for marker in _CJK_FONT_MARKERS
    )


def _find_executable(
    names: tuple[str, ...],
    *,
    macos_app: Path | None = None,
) -> str:
    for name in names:
        resolved = shutil.which(name)
        if resolved:
            return resolved
    if macos_app is not None and macos_app.is_file():
        return str(macos_app)
    raise ExportError(f"DOCX visual QA dependency not found: {' or '.join(names)}")


def _macos_fontconfig(path: Path) -> None:
    user_fonts = Path.home() / "Library" / "Fonts"
    path.write_text(
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE fontconfig SYSTEM "fonts.dtd">\n'
        "<fontconfig>\n"
        "  <dir>/System/Library/Fonts</dir>\n"
        "  <dir>/Library/Fonts</dir>\n"
        f"  <dir>{escape(str(user_fonts))}</dir>\n"
        f"  <cachedir>{path.parent / 'fontconfig-cache'}</cachedir>\n"
        "</fontconfig>\n",
        encoding="utf-8",
    )


def _run(
    command: list[str],
    *,
    env: dict[str, str],
    timeout: float,
) -> subprocess.CompletedProcess[str]:
    try:
        return subprocess.run(
            command,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
            env=env,
            timeout=timeout,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise ExportError(
            f"DOCX visual QA command failed: {type(exc).__name__}"
        ) from exc


def _render_docx(
    docx_path: Path,
    output_dir: Path,
    *,
    timeout: float,
) -> tuple[Path, list[Path], dict[str, Any]]:
    soffice = _find_executable(
        ("soffice", "libreoffice"),
        macos_app=Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    )
    pdftoppm = _find_executable(("pdftoppm",))
    with tempfile.TemporaryDirectory(prefix="patent-docx-qa-") as temp_name:
        temp_dir = Path(temp_name)
        profile = temp_dir / "lo-profile"
        converted = temp_dir / "converted"
        profile.mkdir()
        converted.mkdir()
        env = os.environ.copy()
        env["HOME"] = str(profile)
        env["TMPDIR"] = "/private/tmp" if sys.platform == "darwin" else temp_name
        if sys.platform == "darwin":
            fontconfig = temp_dir / "fonts.conf"
            _macos_fontconfig(fontconfig)
            env["FONTCONFIG_FILE"] = str(fontconfig)
        command = [
            soffice,
            f"-env:UserInstallation=file://{profile / 'user'}",
            "--invisible",
            "--headless",
            "--norestore",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(converted),
            str(docx_path),
        ]
        result = _run(command, env=env, timeout=timeout)
        emitted = sorted(converted.glob("*.pdf"))
        if (
            result.returncode != 0
            or len(emitted) != 1
            or emitted[0].stat().st_size == 0
        ):
            detail = (
                result.stderr or result.stdout or "no PDF emitted"
            ).strip()[-1000:]
            raise ExportError(
                f"DOCX render failed with exit {result.returncode}: {detail}"
            )
        pdf_path = output_dir / "rendered.pdf"
        shutil.copy2(emitted[0], pdf_path)

    prefix = output_dir / "page"
    raster = _run(
        [pdftoppm, "-png", "-r", "144", str(pdf_path), str(prefix)],
        env=os.environ.copy(),
        timeout=timeout,
    )
    pages = sorted(
        output_dir.glob("page-*.png"),
        key=lambda path: int(path.stem.rsplit("-", 1)[1]),
    )
    if raster.returncode != 0 or not pages:
        detail = (
            raster.stderr or raster.stdout or "no PNG pages emitted"
        ).strip()[-1000:]
        raise ExportError(
            f"PDF rasterization failed with exit {raster.returncode}: {detail}"
        )
    return pdf_path, pages, {
        "soffice": soffice,
        "pdftoppm": pdftoppm,
        "dpi": 144,
    }


def _page_metrics(path: Path) -> dict[str, Any]:
    with Image.open(path) as image:
        gray = image.convert("L")
        width, height = gray.size
        pixels = gray.tobytes()
    ink = sum(1 for value in pixels if value < 245)
    dark = sum(1 for value in pixels if value < 96)
    mask = Image.new("L", (width, height), 0)
    mask.putdata([255 if value < 245 else 0 for value in pixels])
    bbox = mask.getbbox()
    edge_margin = max(8, round(min(width, height) * 0.01))
    touches_edge = bool(
        bbox
        and (
            bbox[0] <= edge_margin
            or bbox[1] <= edge_margin
            or bbox[2] >= width - edge_margin
            or bbox[3] >= height - edge_margin
        )
    )
    return {
        "width_px": width,
        "height_px": height,
        "ink_ratio": round(ink / max(1, width * height), 6),
        "dark_pixel_ratio": round(dark / max(1, width * height), 6),
        "content_bbox": list(bbox) if bbox else None,
        "blank": ink < 100,
        "touches_page_edge": touches_edge,
        "sha256": _sha256(path),
    }


def _pdf_evidence(pdf_path: Path) -> tuple[list[str], str, int]:
    reader = PdfReader(str(pdf_path))
    fonts: set[str] = set()
    texts: list[str] = []
    for page in reader.pages:
        texts.append(page.extract_text() or "")
        resources = page.get("/Resources")
        if resources is None:
            continue
        resources = resources.get_object()
        font_map = resources.get("/Font")
        if font_map is None:
            continue
        for font_ref in font_map.get_object().values():
            base_font = font_ref.get_object().get("/BaseFont")
            if base_font:
                fonts.add(str(base_font).lstrip("/"))
    return sorted(fonts), "\n".join(texts), len(reader.pages)


def _validate_review(
    review_path: Path,
    *,
    docx_sha256: str,
    page_count: int,
) -> dict[str, Any]:
    try:
        review = json.loads(review_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ParseError(
            f"invalid DOCX visual review JSON: {type(exc).__name__}"
        ) from exc
    if review.get("schema_version") != "docx_visual_review_v1":
        raise ParseError(
            "DOCX visual review schema_version must be docx_visual_review_v1"
        )
    if review.get("docx_sha256") != docx_sha256:
        raise ParseError(
            "DOCX visual review docx_sha256 does not match the rendered document"
        )
    if review.get("reviewer_type") not in {"gpt_visual", "human_visual"}:
        raise ParseError(
            "DOCX visual review reviewer_type must be gpt_visual or human_visual"
        )
    pages = review.get("pages")
    if not isinstance(pages, list):
        raise ParseError("DOCX visual review pages must be a list")
    actual_numbers = [
        row.get("page_number") for row in pages if isinstance(row, dict)
    ]
    if actual_numbers != list(range(1, page_count + 1)):
        raise ParseError(
            "DOCX visual review must cover every page once and in order"
        )
    page_verdicts = [row.get("verdict") for row in pages]
    if any(verdict not in {"pass", "rework"} for verdict in page_verdicts):
        raise ParseError(
            "DOCX visual review page verdict must be pass or rework"
        )
    expected_overall = (
        "pass"
        if all(verdict == "pass" for verdict in page_verdicts)
        else "rework"
    )
    if review.get("overall_verdict") != expected_overall:
        raise ParseError(
            "DOCX visual review overall_verdict conflicts with page verdicts"
        )
    return {
        "schema_version": review["schema_version"],
        "reviewer_type": review["reviewer_type"],
        "review_sha256": _sha256(review_path),
        "overall_verdict": expected_overall,
        "pages": pages,
        "notes": str(review.get("notes") or ""),
    }


def default_docx_qa_output(
    outputs_dir: Path,
    docx_path: Path,
) -> Path:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")
    return outputs_dir / "docx_qa" / f"{stamp}-{_sha256(docx_path)[:8]}"


def run_docx_qa(
    docx_path: Path,
    *,
    output_dir: Path,
    allowed_output_root: Path,
    review_path: Path | None = None,
    timeout: float = 120,
) -> dict[str, Any]:
    source = docx_path.expanduser().resolve()
    target = output_dir.expanduser().resolve()
    output_root = allowed_output_root.expanduser().resolve()
    if not source.is_file() or source.suffix.casefold() != ".docx":
        raise ParseError("qa-docx requires an existing .docx file")
    try:
        target.relative_to(output_root)
    except ValueError as exc:
        raise ParseError(
            "qa-docx output directory must stay inside configured outputs_dir"
        ) from exc
    if target == output_root:
        raise ParseError(
            "qa-docx output directory must be a child of configured outputs_dir"
        )
    if target.exists():
        raise ParseError(
            "qa-docx output directory already exists; choose a new directory"
        )
    target.mkdir(parents=True)

    try:
        with ZipFile(source) as archive:
            bad_member = archive.testzip()
        if bad_member:
            raise ParseError(
                f"DOCX ZIP CRC failed at member: {bad_member}"
            )
    except BadZipFile as exc:
        raise ParseError(
            "DOCX is not a valid OOXML ZIP container"
        ) from exc
    document = Document(str(source))
    source_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            paragraph.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        ]
    )
    source_cjk_count = len(_CJK_RE.findall(source_text))
    source_sha256 = _sha256(source)
    pdf_path, pages, renderer = _render_docx(
        source,
        target,
        timeout=timeout,
    )
    font_names, rendered_text, pdf_page_count = _pdf_evidence(pdf_path)
    rendered_cjk_count = len(_CJK_RE.findall(rendered_text))
    metrics = [_page_metrics(page) for page in pages]

    checks = [
        {
            "id": "docx_openable",
            "status": "passed",
            "detail": {
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
            },
        },
        {
            "id": "page_count_consistent",
            "status": (
                "passed"
                if pdf_page_count == len(pages)
                else "failed"
            ),
            "detail": {
                "pdf_pages": pdf_page_count,
                "png_pages": len(pages),
            },
        },
        {
            "id": "no_blank_pages",
            "status": (
                "passed"
                if not any(row["blank"] for row in metrics)
                else "failed"
            ),
            "detail": {
                "blank_pages": [
                    index
                    for index, row in enumerate(metrics, 1)
                    if row["blank"]
                ]
            },
        },
        {
            "id": "no_page_edge_overflow",
            "status": (
                "passed"
                if not any(row["touches_page_edge"] for row in metrics)
                else "failed"
            ),
            "detail": {
                "pages_touching_edge": [
                    index
                    for index, row in enumerate(metrics, 1)
                    if row["touches_page_edge"]
                ]
            },
        },
    ]
    if source_cjk_count:
        ratio = rendered_cjk_count / source_cjk_count
        checks.extend(
            [
                {
                    "id": "cjk_text_preserved",
                    "status": "passed" if ratio >= 0.98 else "failed",
                    "detail": {
                        "source_cjk_characters": source_cjk_count,
                        "rendered_cjk_characters": rendered_cjk_count,
                        "retention_ratio": round(ratio, 6),
                    },
                },
                {
                    "id": "cjk_render_font_evidence",
                    "status": (
                        "passed"
                        if _has_cjk_font_evidence(font_names)
                        else "failed"
                    ),
                    "detail": {"pdf_fonts": font_names},
                },
            ]
        )
    automatic_passed = all(
        row["status"] == "passed" for row in checks
    )
    review = (
        _validate_review(
            review_path.expanduser().resolve(),
            docx_sha256=source_sha256,
            page_count=len(pages),
        )
        if review_path is not None
        else None
    )
    visual_status = (
        "manual_review_required"
        if review is None
        else "passed"
        if review["overall_verdict"] == "pass"
        else "rework"
    )
    overall_status = (
        "failed"
        if not automatic_passed
        else "passed"
        if visual_status == "passed"
        else "rework"
        if visual_status == "rework"
        else "manual_review_required"
    )
    report = {
        "schema_version": "docx_visual_qa_v1",
        "status": overall_status,
        "scope": "document_rendering_and_layout_only",
        "docx": {
            "path": str(source),
            "sha256": source_sha256,
            "contains_cjk": bool(source_cjk_count),
        },
        "renderer": renderer,
        "render_check_status": "passed",
        "automated_layout_status": (
            "passed" if automatic_passed else "failed"
        ),
        "visual_qa_status": visual_status,
        "checks": checks,
        "pages": [
            {
                "page_number": index,
                "path": page.name,
                **row,
            }
            for index, (page, row) in enumerate(
                zip(pages, metrics, strict=True),
                1,
            )
        ],
        "rendered_pdf": {
            "path": pdf_path.name,
            "sha256": _sha256(pdf_path),
        },
        "visual_review": review,
        "truth_boundary": (
            "This report checks rendering and layout only. It is not "
            "patent-professional review, a novelty or inventiveness "
            "conclusion, or legal advice."
        ),
    }
    report_path = target / "report.json"
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    report["report_path"] = str(report_path)
    return report
