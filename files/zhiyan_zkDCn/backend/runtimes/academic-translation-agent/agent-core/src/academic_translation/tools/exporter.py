from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from academic_translation.schemas.models import TranslationRequest, TranslationSegment
from academic_translation.utils.text import restore_protected_content


def source(segment: TranslationSegment) -> str:
    return restore_protected_content(segment.source_text, segment.tokens)


def translation(segment: TranslationSegment) -> str:
    return restore_protected_content(segment.translated_text, segment.tokens)


def set_cjk_font(paragraph) -> None:
    for run in paragraph.runs:
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:eastAsia"), "PingFang SC")
        fonts.set(qn("w:hint"), "eastAsia")


def restore_docx(source_path: Path, output: Path, segments: list[TranslationSegment]) -> None:
    document, by_id = Document(source_path), {segment.segment_id: segment for segment in segments}
    for index, paragraph in enumerate(document.paragraphs, start=1):
        segment = by_id.get(f"docx-p{index}")
        if not segment:
            continue
        paragraph.clear()
        paragraph.add_run(translation(segment) if segment.translatable else source(segment))
        set_cjk_font(paragraph)
    for table_index, table in enumerate(document.tables, start=1):
        seen: set[int] = set()
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                segment = by_id.get(f"table{table_index}-r{row_index}-c{cell_index}")
                if not segment:
                    continue
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                paragraph.add_run(translation(segment) if segment.translatable else source(segment))
                set_cjk_font(paragraph)
    for section_index, section in enumerate(document.sections, start=1):
        for part_name, container in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
                segment = by_id.get(f"{part_name}-{section_index}-p{paragraph_index}")
                if not segment or not segment.translatable:
                    continue
                paragraph.clear()
                paragraph.add_run(translation(segment))
                set_cjk_font(paragraph)
    document.save(output)


def export_translation(request: TranslationRequest, segments: list[TranslationSegment], glossary: dict[str, str], quality: dict) -> dict[str, str]:
    output_dir = request.output_dir or Path("outputs")
    output_dir.mkdir(parents=True, exist_ok=True)
    base = request.input_path.stem if request.input_path else "selection"
    mono, bilingual = output_dir / f"{base}-{request.target_lang}.md", output_dir / f"{base}-bilingual.md"
    report, docx = output_dir / f"{base}-translation-report.json", output_dir / f"{base}-{request.target_lang}.docx"
    mono.write_text("\n\n".join(translation(item) for item in segments if item.translatable), encoding="utf-8")
    if request.include_bilingual_markdown:
        bilingual.write_text("\n\n".join(f"### {item.segment_id}\n\n**Source**\n{source(item)}\n\n**Translation**\n{translation(item) if item.translatable else source(item)}" for item in segments), encoding="utf-8")
    report.write_text(json.dumps({"glossary": glossary, "quality": quality}, ensure_ascii=False, indent=2), encoding="utf-8")
    if request.input_path and request.input_path.suffix.lower() == ".docx":
        restore_docx(request.input_path, docx, segments)
    else:
        document = Document()
        for item in segments:
            value = translation(item) if item.translatable else source(item)
            paragraph = document.add_heading(value, level=1) if item.kind in {"title", "heading"} else document.add_paragraph(value)
            set_cjk_font(paragraph)
        document.save(docx)
    outputs = {"monolingual_markdown": str(mono), "translation_report": str(report), "monolingual_docx": str(docx)}
    if request.include_bilingual_markdown:
        outputs["bilingual_markdown"] = str(bilingual)
    return outputs
