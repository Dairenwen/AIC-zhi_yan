from pathlib import Path

from docx import Document

from academic_translation.schemas.models import ElementPolicy
from academic_translation.tools.document_parser import parse_document


def test_parser_protects_formula_and_references() -> None:
    segments = parse_document(text="Introduction\n\nWe define $x=y$ [1].\n\nReferences\n\n[1] Paper.", input_path=None, policy=ElementPolicy())
    assert "[[KEEP_0]]" in segments[1].source_text
    assert segments[-1].kind == "reference"
    assert not segments[-1].translatable


def test_docx_parser_marks_headers_footers_and_honors_formula_policy(tmp_path: Path) -> None:
    path = tmp_path / "source.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Header"
    document.sections[0].footer.paragraphs[0].text = "Footer"
    document.add_paragraph("Formula $x=y$.")
    document.save(path)
    preserved = parse_document(text=None, input_path=path, policy=ElementPolicy())
    assert any(item.kind == "header" and not item.translatable for item in preserved)
    assert any(item.kind == "footer" and not item.translatable for item in preserved)
    assert "[[KEEP_0]]" in next(item for item in preserved if item.segment_id == "docx-p1").source_text
    unprotected = parse_document(text=None, input_path=path, policy=ElementPolicy(preserve_formulas=False))
    assert "[[KEEP_0]]" not in next(item for item in unprotected if item.segment_id == "docx-p1").source_text


def test_figure_caption_policy_controls_translation() -> None:
    preserved = parse_document(text="Figure 1: Molecular graph.", input_path=None, policy=ElementPolicy())
    translated = parse_document(text="Figure 1: Molecular graph.", input_path=None, policy=ElementPolicy(preserve_figures=False))
    assert not preserved[0].translatable
    assert translated[0].translatable
