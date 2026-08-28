import re
import threading
import time
from pathlib import Path

from docx import Document

from academic_translation.agent.service import AcademicTranslationAgent
import academic_translation.agent.service as service_module
from academic_translation.schemas.models import TermEntry, TranslationRequest
from academic_translation.tools.terminology_store import TerminologyStore


class FakeModel:
    def generate(self, prompt: str) -> str:
        if "terminology curator" in prompt:
            return '[{"source":"large language model","target":"大语言模型","confidence":0.99}]'
        if "academic copy editor" in prompt:
            tokens = " ".join(re.findall(r"\[\[KEEP_\d+\]\]", prompt))
            return f"润色后的学术译文 {tokens}".strip()
        tokens = " ".join(re.findall(r"\[\[KEEP_\d+\]\]", prompt))
        return f"大语言模型译文 {tokens}".strip()


class HallucinatingModel(FakeModel):
    def generate(self, prompt: str) -> str:
        if "terminology curator" in prompt:
            return "[]"
        return "译文 [[KEEP_999]]"


def test_text_translation_exports_and_checks_tokens(tmp_path: Path) -> None:
    result = AcademicTranslationAgent(FakeModel()).translate(TranslationRequest(text="A large language model uses $x=y$ [1].", glossary={"large language model": "大语言模型"}, output_dir=tmp_path))
    assert result.quality.passed
    assert Path(result.outputs["monolingual_markdown"]).exists()
    assert "bilingual_markdown" not in result.outputs
    assert Path(result.outputs["monolingual_docx"]).exists()


def test_hallucinated_keep_token_is_removed(tmp_path: Path) -> None:
    result = AcademicTranslationAgent(HallucinatingModel()).translate(TranslationRequest(text="A sentence.", output_dir=tmp_path))
    assert "[[KEEP_999]]" not in result.segments[0].translated_text
    assert result.quality.passed


def test_bilingual_markdown_is_opt_in(tmp_path: Path) -> None:
    result = AcademicTranslationAgent(FakeModel()).translate(
        TranslationRequest(text="A sentence.", output_dir=tmp_path, include_bilingual_markdown=True)
    )
    assert Path(result.outputs["bilingual_markdown"]).exists()


def test_docx_layout_assets_are_preserved(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Header preserved"
    heading = document.add_heading("Introduction", level=1)
    document.add_paragraph("A large language model uses $x=y$ [1].")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Method"
    document.save(source)
    result = AcademicTranslationAgent(FakeModel()).translate(TranslationRequest(input_path=source, glossary={"large language model": "大语言模型"}, output_dir=tmp_path / "out"))
    restored = Document(result.outputs["monolingual_docx"])
    assert restored.paragraphs[0].style.name == heading.style.name
    assert "大语言模型译文" in restored.paragraphs[1].text
    assert "$x=y$ [1]" in restored.paragraphs[1].text
    assert "大语言模型译文" in restored.tables[0].cell(0, 0).text
    assert restored.sections[0].header.paragraphs[0].text == "Header preserved"


def test_submission_level_uses_same_model_for_post_editing(tmp_path: Path) -> None:
    result = AcademicTranslationAgent(FakeModel()).translate(
        TranslationRequest(text="A large language model uses $x=y$.", precision="submission", glossary={"large language model": "大语言模型"}, output_dir=tmp_path)
    )
    assert "润色后的学术译文" in result.segments[0].translated_text
    assert "[[KEEP_0]]" in result.segments[0].translated_text


class ParallelModel(FakeModel):
    def __init__(self) -> None:
        self.workers: set[str] = set()

    def generate(self, prompt: str) -> str:
        if "terminology curator" in prompt:
            return "[]"
        self.workers.add(threading.current_thread().name)
        time.sleep(0.03)
        return "译文"


def test_two_segment_parallel_translation_uses_two_workers(tmp_path: Path) -> None:
    model = ParallelModel()
    result = AcademicTranslationAgent(model).translate(
        TranslationRequest(text="First paragraph.\n\nSecond paragraph.", output_dir=tmp_path, max_parallel_segments=2)
    )
    assert result.quality.passed
    assert len(model.workers) >= 2


def test_manual_segment_revision_rechecks_and_reexports(tmp_path: Path) -> None:
    agent = AcademicTranslationAgent(FakeModel())
    request = TranslationRequest(text="A large language model.", glossary={"large language model": "大语言模型"}, output_dir=tmp_path)
    result = agent.translate(request)
    revised = agent.revise_segment(request, result, "s1", "人工修订：大语言模型。")
    assert revised.quality.passed
    assert Path(revised.outputs["monolingual_docx"]).exists()


def test_local_terminology_library_persists_terms_without_other_agents(tmp_path: Path) -> None:
    store = TerminologyStore(tmp_path / "terms.json")
    store.save("materials", [TermEntry(source="molecular relation", target="分子关系", confidence=0.9, origin="model")])
    loaded = store.load("materials")
    assert loaded[0].source == "molecular relation"
    assert loaded[0].target == "分子关系"


def test_pdf_only_requires_a_layout_preserved_pdf(tmp_path: Path) -> None:
    agent = AcademicTranslationAgent(FakeModel())
    try:
        agent.translate(TranslationRequest(text="A sentence.", pdf_only=True, output_dir=tmp_path))
    except ValueError as error:
        assert "--pdf-only" in str(error)
    else:
        raise AssertionError("pdf-only text input should be rejected")


def test_pdf_only_bypasses_segment_graph_and_uses_layout_renderer(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    import fitz
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(source)
    pdf.close()
    captured: dict[str, object] = {}

    def fake_render(request, glossary):
        captured["request"] = request
        captured["glossary"] = glossary
        return {"pdf_monolingual": str(tmp_path / "paper-mono.pdf"), "pdf_layout_pages": "1"}

    monkeypatch.setattr(service_module, "render_pdf_with_pdf2zh", fake_render)
    result = AcademicTranslationAgent(FakeModel()).translate(
        TranslationRequest(input_path=source, preserve_pdf_layout=True, pdf_only=True, glossary={"Method": "方法"}, output_dir=tmp_path)
    )
    assert result.segments == []
    assert result.quality.passed
    assert captured["glossary"]["Method"] == "方法"
    assert result.outputs["pdf_monolingual"].endswith("paper-mono.pdf")
