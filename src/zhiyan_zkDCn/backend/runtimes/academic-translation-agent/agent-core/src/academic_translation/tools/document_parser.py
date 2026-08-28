from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document

from academic_translation.schemas.models import ElementPolicy, TranslationSegment
from academic_translation.utils.text import mask_protected_content, normalize_whitespace


REFERENCE = re.compile(r"^(references|bibliography|参考文献)\s*$", re.I)
HEADING = re.compile(r"^(\d+(?:\.\d+)*\.?\s+.+|[A-Z][A-Z\s]{4,})$")


def classify(text: str, in_references: bool) -> tuple[str, bool]:
    if REFERENCE.match(text):
        return "heading", True
    if in_references:
        return "reference", True
    if HEADING.match(text) or (len(text) < 90 and text.endswith(":")):
        return "heading", False
    if re.match(r"^(figure|fig\.|table|图|表)\s*\d+", text, re.I):
        return "caption", False
    return "paragraph", False


def make_segment(segment_id: str, text: str, policy: ElementPolicy, *, kind: str = "paragraph", page: int | None = None, style: str | None = None, translatable: bool | None = None) -> TranslationSegment:
    clean = normalize_whitespace(text)
    source, tokens = mask_protected_content(clean) if policy.preserve_formulas else (clean, {})
    return TranslationSegment(
        segment_id=segment_id,
        kind=kind,  # type: ignore[arg-type]
        source_text=source,
        translatable=translatable if translatable is not None else not (
            (kind == "reference" and policy.preserve_references)
            or (kind == "caption" and policy.preserve_figures)
        ),
        page=page,
        style=style,
        tokens=tokens,
    )


def parse_text(text: str, policy: ElementPolicy) -> list[TranslationSegment]:
    segments: list[TranslationSegment] = []
    in_references = False
    for raw in re.split(r"\n\s*\n", text):
        clean = normalize_whitespace(raw)
        if not clean:
            continue
        kind, starts_references = classify(clean, in_references)
        in_references = in_references or starts_references
        segments.append(make_segment(f"s{len(segments)+1}", clean, policy, kind=kind))
    return segments


def parse_docx(path: Path, policy: ElementPolicy) -> list[TranslationSegment]:
    document, segments, in_references = Document(path), [], False
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_whitespace(paragraph.text)
        if not text:
            continue
        kind, starts_references = classify(text, in_references)
        in_references = in_references or starts_references
        segments.append(make_segment(f"docx-p{index}", text, policy, kind=kind, style=paragraph.style.name))
    for table_index, table in enumerate(document.tables, start=1):
        seen: set[int] = set()
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                if normalize_whitespace(cell.text):
                    segments.append(make_segment(f"table{table_index}-r{row_index}-c{cell_index}", cell.text, policy, kind="table"))
    for section_index, section in enumerate(document.sections, start=1):
        for part_name, container in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
                text = normalize_whitespace(paragraph.text)
                if text:
                    segments.append(make_segment(f"{part_name}-{section_index}-p{paragraph_index}", text, policy, kind=part_name, translatable=not policy.preserve_headers_footers))
    return segments


def parse_pdf(path: Path, policy: ElementPolicy) -> list[TranslationSegment]:
    document, segments, in_references = fitz.open(path), [], False
    for page_number, page in enumerate(document, start=1):
        for block in page.get_text("blocks", sort=True):
            text = normalize_whitespace(block[4])
            if not text:
                continue
            kind, starts_references = classify(text, in_references)
            in_references = in_references or starts_references
            segments.append(make_segment(f"pdf-p{page_number}-s{len(segments)+1}", text, policy, kind=kind, page=page_number))
    document.close()
    return segments


def parse_document(*, text: str | None, input_path: Path | None, policy: ElementPolicy) -> list[TranslationSegment]:
    if text is not None:
        return parse_text(text, policy)
    if input_path is None:
        raise ValueError("No input was provided.")
    if input_path.suffix.lower() in {".txt", ".md"}:
        return parse_text(input_path.read_text(encoding="utf-8"), policy)
    if input_path.suffix.lower() == ".docx":
        return parse_docx(input_path, policy)
    if input_path.suffix.lower() == ".pdf":
        return parse_pdf(input_path, policy)
    raise ValueError("Supported inputs: .txt, .md, .docx, .pdf")
