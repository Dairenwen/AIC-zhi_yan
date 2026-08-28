from __future__ import annotations

import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from patent_agent.errors import ParseError
from patent_agent.platform_support import font_available


@dataclass(frozen=True)
class DocxFontResolution:
    configured_font: str
    effective_font: str
    configured_font_available: bool
    fallback_used: bool


def resolve_docx_font(configured_font: str) -> DocxFontResolution:
    """Resolve one installed CJK font without installing or bundling font files."""
    configured = configured_font.strip() or "Noto Sans CJK SC"
    configured_available, _detector = font_available(configured)
    if configured_available:
        return DocxFontResolution(configured, configured, True, False)

    if sys.platform == "darwin":
        fallbacks = ("PingFang SC", "Songti SC", "Arial Unicode MS")
    elif os.name == "nt":
        fallbacks = ("Microsoft YaHei", "SimSun", "Arial Unicode MS")
    else:
        fallbacks = (
            "Noto Sans CJK SC",
            "Source Han Sans SC",
            "WenQuanYi Zen Hei",
            "Arial Unicode MS",
        )
    for candidate in fallbacks:
        available, _detector = font_available(candidate)
        if available:
            return DocxFontResolution(configured, candidate, False, True)
    return DocxFontResolution(configured, configured, False, False)


def _expand_single_line_display_math(markdown: str) -> str:
    """Make vendored md_to_docx recognize already-rendered single-line display math."""
    return re.sub(r"^\\\[(.+)\\\]\s*$", lambda m: "\\[\n" + m.group(1) + "\n\\]", markdown, flags=re.MULTILINE)


def _normalize_inline_math_delimiters(markdown: str) -> str:
    """Use dollar delimiters so upstream table parsing handles nested parentheses."""
    return re.sub(r"\\\((.*?)\\\)", lambda m: "$" + m.group(1) + "$", markdown, flags=re.DOTALL)


def _quote_mermaid_node_labels(markdown: str) -> str:
    """Quote Mermaid node labels so formulas and punctuation remain parseable."""

    def normalize_block(match: re.Match[str]) -> str:
        body = match.group(1)

        def quote_square(node: re.Match[str]) -> str:
            raw_label = node.group(2).strip()
            if len(raw_label) >= 2 and raw_label[0] == raw_label[-1] and raw_label[0] in {'"', "'"}:
                return node.group(0)
            label = raw_label.replace('"', "'")
            return f'{node.group(1)}["{label}"]'

        def quote_diamond(node: re.Match[str]) -> str:
            raw_label = node.group(2).strip()
            if len(raw_label) >= 2 and raw_label[0] == raw_label[-1] and raw_label[0] in {'"', "'"}:
                return node.group(0)
            label = raw_label.replace('"', "'")
            return f'{node.group(1)}{{"{label}"}}'

        body = re.sub(r"\b([A-Za-z][A-Za-z0-9_]*)\[([^\]\n]+)\]", quote_square, body)
        body = re.sub(r"\b([A-Za-z][A-Za-z0-9_]*)\{([^}\n]+)\}", quote_diamond, body)
        return "```mermaid\n" + body + "```"

    return re.sub(r"```mermaid\s*\n(.*?)```", normalize_block, markdown, flags=re.DOTALL | re.IGNORECASE)


def _default_cjk_font(configured_font: str | None = None) -> str:
    configured = (
        configured_font
        if configured_font is not None
        else os.environ.get("PATENT_AGENT_DOCX_FONT", "")
    ).strip()
    return resolve_docx_font(configured or "Noto Sans CJK SC").effective_font


def is_font_available(font_name: str) -> bool:
    available, _detector = font_available(font_name)
    return available


def _set_font(run_or_style, font_name: str, size_pt: float | None = None) -> None:
    font = run_or_style.font
    font.name = font_name
    if size_pt is not None:
        font.size = Pt(size_pt)
    rpr = run_or_style._element.get_or_add_rPr() if hasattr(run_or_style._element, "get_or_add_rPr") else run_or_style._element.rPr
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), font_name)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    cols = max(len(row.cells) for row in table.rows)
    base = width_dxa // cols
    widths = [base] * cols
    widths[-1] += width_dxa - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_idx == 0:
            if tr_pr.find(qn("w:tblHeader")) is None:
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def apply_standard_business_brief(
    docx_path: Path,
    configured_font: str | None = None,
) -> None:
    """Apply explicit standard_business_brief tokens and CJK-safe fonts."""
    font_name = _default_cjk_font(configured_font)
    doc = Document(str(docx_path))
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    style_tokens = {
        "Normal": (11, 0, 6, 1.10, None),
        "Heading 1": (16, 16, 8, 1.0, "2E74B5"),
        "Heading 2": (13, 12, 6, 1.0, "2E74B5"),
        "Heading 3": (12, 8, 4, 1.0, "1F4D78"),
        "List Bullet": (11, 0, 8, 1.167, None),
        "List Number": (11, 0, 8, 1.167, None),
    }
    for name, (size, before, after, line, color) in style_tokens.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        _set_font(style, font_name, size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line
        if color:
            style.font.color.rgb = RGBColor.from_string(color)

    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        for run in paragraph.runs:
            _set_font(run, font_name, run.font.size.pt if run.font.size else None)
    for table in doc.tables:
        _set_table_geometry(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.10
                    for run in paragraph.runs:
                        _set_font(run, font_name, run.font.size.pt if run.font.size else 10)
    doc.save(str(docx_path))


def export_markdown_and_docx(
    *,
    vendor_root: Path,
    checked_md: Path,
    final_md: Path,
    final_docx: Path,
    environment: dict[str, str],
    timeout_seconds: float = 600,
) -> tuple[str, str]:
    mermaid = vendor_root / "tools" / "mermaid_render.py"
    export_environment = dict(environment)
    export_environment["PUPPETEER_CACHE_DIR"] = str(vendor_root / "tools" / ".puppeteer-cache")
    export_environment["PYTHONIOENCODING"] = "utf-8"
    export_environment["PYTHONUTF8"] = "1"
    normalized = _normalize_inline_math_delimiters(checked_md.read_text(encoding="utf-8"))
    normalized = _quote_mermaid_node_labels(normalized)
    with tempfile.TemporaryDirectory(prefix="patent-export-") as source_tmpdir:
        normalized_source = Path(source_tmpdir) / "normalized.md"
        normalized_source.write_text(normalized, encoding="utf-8")
        render_proc = subprocess.run(
            [sys.executable, str(mermaid), "-i", str(normalized_source), "-o", str(final_md), "--no-docx"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout_seconds,
            env=export_environment,
        )
    logs = (render_proc.stdout or "") + ("\n" if render_proc.stdout and render_proc.stderr else "") + (render_proc.stderr or "")
    if render_proc.returncode != 0 or not final_md.is_file():
        raise ParseError("upstream Mermaid/Markdown export failed; details saved to document_export.log")

    docx_script = vendor_root / "tools" / "md_to_docx.py"
    markdown = _expand_single_line_display_math(final_md.read_text(encoding="utf-8"))
    with tempfile.TemporaryDirectory(prefix="patent-docx-") as tmpdir:
        temp_md = Path(tmpdir) / "docx_source.md"
        temp_md.write_text(markdown, encoding="utf-8")
        docx_proc = subprocess.run(
            [
                sys.executable,
                str(docx_script),
                "-i",
                str(temp_md),
                "-o",
                str(final_docx),
                "--base-dir",
                str(final_md.parent),
                "--no-math-render",
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout_seconds,
            env=export_environment,
        )
    logs += "\n" + (docx_proc.stdout or "") + ("\n" if docx_proc.stdout and docx_proc.stderr else "") + (docx_proc.stderr or "")
    if docx_proc.returncode != 0 or not final_docx.is_file():
        return logs, "failed"
    apply_standard_business_brief(
        final_docx,
        export_environment.get("PATENT_AGENT_DOCX_FONT"),
    )
    return logs, "passed"
